"""
Word document generator for meeting summaries and intelligence signals.

Generates .docx files from structured data using python-docx.
Word documents are more shareable with external stakeholders than Markdown.

Usage:
    from services.word_generator import generate_summary_docx, generate_signal_docx

    docx_bytes = generate_summary_docx(meeting_title="Strategy Meeting", ...)
    docx_bytes = generate_signal_docx(signal_content="## FLAGS\n...", week_number=14, year=2026)
"""

import io
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def generate_summary_docx(
    meeting_title: str,
    meeting_date: str,
    participants: list[str],
    duration_minutes: int,
    sensitivity: str,
    decisions: list[dict],
    tasks: list[dict],
    follow_ups: list[dict],
    open_questions: list[dict],
    discussion_summary: str,
    stakeholders_mentioned: list[dict] | None = None,
) -> bytes:
    """
    Generate a Word document (.docx) from structured meeting data.

    Args:
        meeting_title: Title of the meeting.
        meeting_date: Date string (e.g. "2026-03-01").
        participants: List of participant names.
        duration_minutes: Meeting duration in minutes.
        sensitivity: 'normal' or 'sensitive'.
        decisions: List of decision dicts with 'description', 'participants_involved', 'transcript_timestamp'.
        tasks: List of task dicts with 'title', 'assignee', 'deadline', 'priority', 'status', 'transcript_timestamp'.
        follow_ups: List of follow-up dicts with 'title', 'led_by', 'proposed_date', 'participants'.
        open_questions: List of open question dicts with 'question', 'raised_by'.
        discussion_summary: Prose summary of the discussion.
        stakeholders_mentioned: Optional list of stakeholder dicts.

    Returns:
        Bytes content of the .docx file.
    """
    doc = Document()

    # --- Title ---
    title_para = doc.add_heading(f"Meeting Summary: {meeting_title}", level=1)

    # --- Metadata ---
    participants_str = ", ".join(participants) if participants else "Not recorded"
    duration_str = f"{duration_minutes} minutes" if duration_minutes else "Not recorded"
    sensitivity_str = sensitivity.capitalize() if sensitivity else "Normal"

    meta_text = (
        f"Date: {meeting_date}\n"
        f"Duration: {duration_str}\n"
        f"Participants: {participants_str}\n"
        f"Sensitivity: {sensitivity_str}"
    )
    meta_para = doc.add_paragraph(meta_text)
    meta_para.style = doc.styles["No Spacing"]

    doc.add_paragraph()  # spacing

    # --- Discussion Summary (at the top for quick scanning) ---
    doc.add_heading("Summary", level=2)
    if discussion_summary:
        # Truncate to ~800 chars, but end at a sentence boundary
        if len(discussion_summary) > 800:
            # Find last period before 800 chars
            cut = discussion_summary[:800].rfind(".")
            if cut > 400:
                summary_text = discussion_summary[:cut + 1]
            else:
                # No good sentence break — cut at last space
                cut = discussion_summary[:800].rfind(" ")
                summary_text = discussion_summary[:cut] + "..." if cut > 0 else discussion_summary[:800] + "..."
        else:
            summary_text = discussion_summary
        para = doc.add_paragraph(summary_text)
        for run in para.runs:
            run.font.size = Pt(10)
    else:
        doc.add_paragraph("No discussion summary available.", style="No Spacing")

    # --- Key Decisions ---
    doc.add_heading("Key Decisions", level=2)
    if decisions:
        for i, d in enumerate(decisions, 1):
            label = d.get("label", "")
            desc = d.get("description", "")
            who = d.get("participants_involved", ["team"])
            who_str = ", ".join(who) if isinstance(who, list) else str(who)
            prefix = f"[{label}] " if label else ""
            text = f"{prefix}{desc} — {who_str}"
            doc.add_paragraph(text, style="List Number")
    else:
        doc.add_paragraph("No key decisions recorded.", style="No Spacing")

    # --- Action Items (compact table) ---
    doc.add_heading("Action Items", level=2)
    if tasks:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        headers = ["Pri", "Action Item", "Owner", "Deadline"]
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.paragraph_format.space_after = Pt(2)

        for t in tasks:
            row = table.add_row()
            row.cells[0].text = t.get("priority", "M")
            # Use label + title for compact display
            label = t.get("label", "")
            title = t.get("title", "")
            row.cells[1].text = f"[{label}] {title}" if label else title
            row.cells[2].text = t.get("assignee", "") or "—"
            row.cells[3].text = str(t.get("deadline", "") or "—")

            # Compact row spacing
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    else:
        doc.add_paragraph("No action items recorded.", style="No Spacing")

    # --- Follow-Up Meetings ---
    doc.add_heading("Follow-Up Meetings", level=2)
    if follow_ups:
        for fu in follow_ups:
            title = fu.get("title", "")
            led_by = fu.get("led_by", "TBD")
            proposed = fu.get("proposed_date", "TBD")
            parts = fu.get("participants", [])
            parts_str = ", ".join(parts) if isinstance(parts, list) else str(parts)
            text = f"{title} — Led by {led_by}"
            if proposed:
                text += f", proposed {proposed}"
            if parts_str:
                text += f" (with {parts_str})"
            doc.add_paragraph(text, style="List Bullet")
    else:
        doc.add_paragraph("No follow-up meetings scheduled.", style="No Spacing")

    # --- Open Questions ---
    doc.add_heading("Open Questions", level=2)
    if open_questions:
        for q in open_questions:
            question = q.get("question", "")
            raised_by = q.get("raised_by", "")
            text = question
            if raised_by:
                text += f" (raised by {raised_by})"
            doc.add_paragraph(text, style="List Bullet")
    else:
        doc.add_paragraph("No open questions.", style="No Spacing")

    # --- Stakeholders ---
    if stakeholders_mentioned:
        doc.add_heading("Stakeholders Mentioned", level=2)
        for s in stakeholders_mentioned:
            name = s.get("name", s.get("organization", ""))
            role = s.get("role", s.get("type", ""))
            text = name
            if role:
                text += f" ({role})"
            doc.add_paragraph(text, style="List Bullet")

    # --- Footer ---
    doc.add_paragraph()  # spacing
    footer_para = doc.add_paragraph("Generated by Gianluigi")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Serialize to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_prep_docx(
    title: str,
    date: str,
    meeting_type: str,
    participants: list[str],
    sections: list[dict],
    gantt_snapshot: list[list[str]] | None = None,
    focus_areas: list[str] | None = None,
) -> bytes:
    """
    Generate a Word document for a meeting prep.

    Args:
        title: Meeting title.
        date: Date string.
        meeting_type: Template display name.
        participants: List of participant names.
        sections: List of section dicts with name, status, data, item_count.
        gantt_snapshot: Optional table rows [section, item, status, owner, week].
        focus_areas: Optional focus instructions from Eyal.

    Returns:
        Bytes content of the .docx file.
    """
    doc = Document()

    # --- Title ---
    doc.add_heading(f"Meeting Prep: {title}", level=1)

    # --- Metadata ---
    participants_str = ", ".join(participants) if participants else "Not recorded"
    meta_text = (
        f"Date: {date}\n"
        f"Meeting Type: {meeting_type}\n"
        f"Participants: {participants_str}"
    )
    meta_para = doc.add_paragraph(meta_text)
    meta_para.style = doc.styles["No Spacing"]
    doc.add_paragraph()

    # --- Focus Areas ---
    if focus_areas:
        doc.add_heading("Focus Areas", level=2)
        for fa in focus_areas:
            doc.add_paragraph(fa, style="List Bullet")

    # --- Sections ---
    for section in sections:
        name = section.get("name", "Section")
        status = section.get("status", "ok")
        data = section.get("data")
        item_count = section.get("item_count", 0)

        doc.add_heading(name, level=2)

        if "unavailable" in status:
            doc.add_paragraph(f"Data unavailable: {status}", style="No Spacing")
            continue

        if not data or item_count == 0:
            doc.add_paragraph("No items found.", style="No Spacing")
            continue

        # Format data based on type
        if isinstance(data, dict):
            for key, items in data.items():
                doc.add_paragraph(str(key), style="Heading 3")
                if isinstance(items, list):
                    for item in items:
                        if isinstance(item, dict):
                            text = item.get("title", item.get("description", str(item)))
                            doc.add_paragraph(str(text), style="List Bullet")
                        else:
                            doc.add_paragraph(str(item), style="List Bullet")
                else:
                    doc.add_paragraph(str(items))
        elif isinstance(data, list):
            for item in data:
                if isinstance(item, dict):
                    text = item.get("title", item.get("description", item.get("commitment", str(item))))
                    doc.add_paragraph(str(text), style="List Bullet")
                else:
                    doc.add_paragraph(str(item), style="List Bullet")

    # --- Gantt Snapshot ---
    if gantt_snapshot:
        doc.add_heading("Gantt Status", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        headers = ["Section", "Item", "Status", "Owner", "Week"]
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for row_data in gantt_snapshot:
            row = table.add_row()
            for j, val in enumerate(row_data[:5]):
                row.cells[j].text = str(val)

    # --- Footer ---
    doc.add_paragraph()
    from datetime import datetime
    footer_para = doc.add_paragraph(
        f"Generated by Gianluigi on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_signal_docx(
    signal_content: str,
    week_number: int,
    year: int,
    flags: list[dict] | None = None,
    research_source: str = "perplexity",
) -> bytes:
    """
    Generate a Word document for an intelligence signal.

    Parses the markdown signal content into formatted sections with
    CropSight branding, flags table, and professional layout.

    Args:
        signal_content: Full signal report in markdown.
        week_number: ISO week number.
        year: Year.
        flags: Optional list of flag dicts [{flag, urgency}].
        research_source: "perplexity", "perplexity_retry", or "claude_search".

    Returns:
        Bytes content of the .docx file.
    """
    import re
    from datetime import datetime

    doc = Document()

    # --- Title ---
    doc.add_heading(
        f"CropSight Intelligence Signal — W{week_number}/{year}", level=1
    )

    # --- Subtitle with accent color ---
    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Weekly Market Intelligence")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0xD4, 0xAA)  # CropSight green
    subtitle.style = doc.styles["No Spacing"]

    # --- Metadata ---
    meta_text = f"Generated: {datetime.now().strftime('%B %d, %Y')}"
    if research_source and research_source != "perplexity":
        meta_text += f"\nResearch source: {research_source} (backup)"
    meta_para = doc.add_paragraph(meta_text)
    meta_para.style = doc.styles["No Spacing"]
    for run in meta_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # spacing

    # --- Flags section ---
    if flags:
        doc.add_heading("Flags", level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"

        for f in flags:
            if not isinstance(f, dict):
                continue
            row = table.add_row()
            urgency = f.get("urgency", "medium")

            urgency_cell = row.cells[0]
            urgency_cell.text = urgency.upper()
            for paragraph in urgency_cell.paragraphs:
                for r in paragraph.runs:
                    r.bold = True
                    r.font.size = Pt(9)
                    if urgency == "high":
                        r.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)
                    else:
                        r.font.color.rgb = RGBColor(0xFF, 0xC1, 0x07)

            flag_cell = row.cells[1]
            flag_cell.text = f.get("flag", "")
            for paragraph in flag_cell.paragraphs:
                for r in paragraph.runs:
                    r.font.size = Pt(10)

        for row in table.rows:
            row.cells[0].width = Inches(0.8)
            row.cells[1].width = Inches(5.5)

        doc.add_paragraph()

    # --- Parse markdown content into sections ---
    sections = _parse_signal_sections(signal_content)

    for section_name, section_body in sections:
        if section_name.upper().startswith("FLAG"):
            continue

        doc.add_heading(section_name, level=2)

        if not section_body.strip():
            doc.add_paragraph("No notable developments this week.", style="No Spacing")
            continue

        lines = section_body.strip().split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith("- ") or stripped.startswith("* "):
                para = doc.add_paragraph(style="List Bullet")
                _add_formatted_runs(para, stripped[2:])
            else:
                para = doc.add_paragraph()
                _add_formatted_runs(para, stripped)
                for r in para.runs:
                    r.font.size = Pt(10)

    # --- Footer ---
    doc.add_paragraph()
    footer_para = doc.add_paragraph(
        f"Generated by Gianluigi — CropSight Intelligence Signal W{week_number}/{year}"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _parse_signal_sections(content: str) -> list[tuple[str, str]]:
    """Parse markdown signal content into (section_name, section_body) tuples."""
    import re

    parts = re.split(r'^(#{2,3})\s+(.+)$', content, flags=re.MULTILINE)

    sections = []
    i = 0
    while i < len(parts):
        if i + 2 < len(parts) and parts[i].startswith("#"):
            section_name = parts[i + 1].strip()
            section_body = parts[i + 2] if i + 2 < len(parts) else ""
            sections.append((section_name, section_body))
            i += 3
        else:
            i += 1

    return sections


def _add_formatted_runs(paragraph, text: str) -> None:
    """Add text with **bold** markdown formatting to a paragraph."""
    import re

    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
