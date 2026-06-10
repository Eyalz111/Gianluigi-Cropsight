"""PR8 — branded summary artifacts (.docx + email) behind SUMMARY_BRANDED_ENABLED.

Off = today's plain document / 4-column email, byte-for-byte. On = CropSight
palette + the Area + Urgency columns the structured outputs drop today.
"""
from unittest.mock import patch

import pytest
from docx import Document

from config.settings import settings
from services.word_generator import generate_summary_docx, BRAND_GREEN


def _sample():
    return dict(
        meeting_title="Strategy Sync",
        meeting_date="2026-06-10",
        participants=["Eyal", "Roye"],
        duration_minutes=40,
        sensitivity="founders",
        decisions=[{"description": "Adopt Postgres", "participants_involved": ["Roye"]}],
        tasks=[{"title": "Ship pilot", "assignee": "Roye", "deadline": "2026-06-20",
                "priority": "H", "urgency": "H", "area_label": "Product & Tech"}],
        follow_ups=[], open_questions=[], discussion_summary="We discussed the stack.",
    )


def _action_items_table(docx_bytes):
    import io
    doc = Document(io.BytesIO(docx_bytes))
    # the Action Items table is the one whose first cell header is "Pri"
    for tbl in doc.tables:
        if tbl.rows[0].cells[0].text.strip() == "Pri":
            return tbl
    return None


class TestDocx:
    def test_off_is_four_column_unbranded(self):
        with patch.object(settings, "SUMMARY_BRANDED_ENABLED", False):
            data = generate_summary_docx(**_sample())
        tbl = _action_items_table(data)
        assert tbl is not None
        assert len(tbl.columns) == 4
        headers = [c.text for c in tbl.rows[0].cells]
        assert headers == ["Pri", "Action Item", "Owner", "Deadline"]

    def test_on_adds_area_and_urgency_columns(self):
        with patch.object(settings, "SUMMARY_BRANDED_ENABLED", True):
            data = generate_summary_docx(**_sample())
        tbl = _action_items_table(data)
        assert len(tbl.columns) == 6
        headers = [c.text for c in tbl.rows[0].cells]
        assert headers == ["Pri", "Action Item", "Area", "Owner", "Deadline", "Urgency"]
        # the data row carries the new fields
        data_row = [c.text for c in tbl.rows[1].cells]
        assert "Product & Tech" in data_row
        assert "H" in data_row

    def test_on_title_is_brand_green(self):
        with patch.object(settings, "SUMMARY_BRANDED_ENABLED", True):
            data = generate_summary_docx(**_sample())
        import io
        doc = Document(io.BytesIO(data))
        title = doc.paragraphs[0]
        assert "Meeting Summary" in title.text
        assert any(r.font.color and r.font.color.rgb == BRAND_GREEN for r in title.runs)

    def test_both_paths_produce_valid_docx(self):
        # a docx that opens without raising is the floor-safety check
        for flag in (False, True):
            with patch.object(settings, "SUMMARY_BRANDED_ENABLED", flag):
                data = generate_summary_docx(**_sample())
            import io
            assert Document(io.BytesIO(data)) is not None
