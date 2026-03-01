"""
Tests for services/word_generator.py

Pure unit tests — no mocks needed, just verify the generated .docx content.
"""

import io
import pytest
from docx import Document


# Shared test data
SAMPLE_DECISIONS = [
    {
        "description": "Move to AWS for cloud hosting",
        "participants_involved": ["Eyal", "Roye"],
        "transcript_timestamp": "12:34",
    },
    {
        "description": "Hire a new data scientist by Q2",
        "participants_involved": ["Eyal"],
        "transcript_timestamp": "23:45",
    },
]

SAMPLE_TASKS = [
    {
        "title": "Draft AWS migration plan",
        "assignee": "Roye",
        "deadline": "2026-03-15",
        "priority": "H",
        "status": "pending",
        "transcript_timestamp": "15:00",
    },
    {
        "title": "Post data scientist job listing",
        "assignee": "Eyal",
        "deadline": "2026-03-10",
        "priority": "M",
        "status": "pending",
        "transcript_timestamp": "25:10",
    },
]

SAMPLE_FOLLOW_UPS = [
    {
        "title": "AWS Architecture Review",
        "led_by": "Roye",
        "proposed_date": "2026-03-20",
        "participants": ["Eyal", "Roye", "Yoram"],
    },
]

SAMPLE_OPEN_QUESTIONS = [
    {"question": "Which AWS region to use?", "raised_by": "Roye"},
    {"question": "Budget for data scientist hire?", "raised_by": "Eyal"},
]

SAMPLE_STAKEHOLDERS = [
    {"name": "Jason Adelman", "role": "Investor"},
    {"organization": "IIA", "type": "Government Agency"},
]


def _generate_sample_docx(**overrides):
    """Helper to generate a docx with sample data, allowing overrides."""
    from services.word_generator import generate_summary_docx

    defaults = {
        "meeting_title": "Strategy Planning",
        "meeting_date": "2026-03-01",
        "participants": ["Eyal", "Roye", "Paolo"],
        "duration_minutes": 60,
        "sensitivity": "normal",
        "decisions": SAMPLE_DECISIONS,
        "tasks": SAMPLE_TASKS,
        "follow_ups": SAMPLE_FOLLOW_UPS,
        "open_questions": SAMPLE_OPEN_QUESTIONS,
        "discussion_summary": "Discussed cloud migration strategy and hiring plans.",
        "stakeholders_mentioned": SAMPLE_STAKEHOLDERS,
    }
    defaults.update(overrides)
    return generate_summary_docx(**defaults)


class TestWordGenerator:
    """Tests for generate_summary_docx."""

    def test_returns_bytes(self):
        """Output is bytes (not str or BytesIO)."""
        result = _generate_sample_docx()
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_valid_docx(self):
        """Output can be parsed as a valid .docx document."""
        result = _generate_sample_docx()
        doc = Document(io.BytesIO(result))
        # Should have at least title + sections
        assert len(doc.paragraphs) > 0

    def test_title_present(self):
        """Meeting title appears in the document heading."""
        result = _generate_sample_docx(meeting_title="Board Review Q1")
        doc = Document(io.BytesIO(result))
        # First paragraph should be the heading
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Board Review Q1" in all_text

    def test_decisions_present(self):
        """Decisions appear as numbered items."""
        result = _generate_sample_docx()
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Move to AWS for cloud hosting" in all_text
        assert "Hire a new data scientist" in all_text

    def test_tasks_table_present(self):
        """Tasks appear in a table with correct columns."""
        result = _generate_sample_docx()
        doc = Document(io.BytesIO(result))
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        # Header row
        headers = [cell.text for cell in table.rows[0].cells]
        assert "#" in headers
        assert "Task" in headers
        assert "Assignee" in headers
        # Data rows
        assert len(table.rows) >= 3  # header + 2 tasks
        assert "Roye" in table.rows[1].cells[2].text

    def test_follow_ups_present(self):
        """Follow-up meetings appear in the document."""
        result = _generate_sample_docx()
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "AWS Architecture Review" in all_text
        assert "Roye" in all_text

    def test_open_questions_present(self):
        """Open questions appear in the document."""
        result = _generate_sample_docx()
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Which AWS region to use?" in all_text
        assert "Budget for data scientist hire?" in all_text

    def test_discussion_summary_present(self):
        """Discussion summary text appears in the document."""
        result = _generate_sample_docx()
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "cloud migration strategy" in all_text

    def test_stakeholders_present(self):
        """Stakeholders section appears when provided."""
        result = _generate_sample_docx()
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jason Adelman" in all_text
        assert "IIA" in all_text

    def test_empty_sections_handled(self):
        """Empty data produces graceful fallback text."""
        result = _generate_sample_docx(
            decisions=[],
            tasks=[],
            follow_ups=[],
            open_questions=[],
            discussion_summary="",
            stakeholders_mentioned=None,
        )
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "No key decisions recorded" in all_text
        assert "No action items recorded" in all_text
        assert "No follow-up meetings scheduled" in all_text
        assert "No open questions" in all_text

    def test_footer_present(self):
        """Footer with 'Generated by Gianluigi' appears."""
        result = _generate_sample_docx()
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Generated by Gianluigi" in all_text

    def test_metadata_present(self):
        """Meeting metadata (date, duration, participants) appears."""
        result = _generate_sample_docx(
            meeting_date="2026-03-01",
            duration_minutes=90,
            participants=["Eyal", "Roye"],
            sensitivity="sensitive",
        )
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "2026-03-01" in all_text
        assert "90 minutes" in all_text
        assert "Eyal, Roye" in all_text
        assert "Sensitive" in all_text
